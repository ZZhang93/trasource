# ============================================================
# core/ingest.py — 多文献类型导入
# 支持：CSV报纸 / PDF（有书签/无书签）/ DOCX / TXT / EPUB / MOBI
# ============================================================

import os
import re
import logging
import chardet
import duckdb
import pandas as pd
from datetime import datetime
from config import CSV_COLUMN_MAP

logger = logging.getLogger(__name__)


# ════════════════════════════════════════════════════════════
# 数据库初始化
# ════════════════════════════════════════════════════════════

def init_database(db_path: str):
    conn = duckdb.connect(db_path)
    try:
        conn.execute("CREATE SEQUENCE IF NOT EXISTS doc_id_seq START 1")
    except Exception:
        pass
    conn.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id                  INTEGER PRIMARY KEY DEFAULT nextval('doc_id_seq'),
            source_file         TEXT,
            file_type           TEXT,
            doc_type            TEXT,
            year                TEXT,
            date                TEXT,
            page                TEXT,
            title               TEXT,
            author              TEXT,
            pub_year            TEXT,
            publisher           TEXT,
            chapter             TEXT,
            section             TEXT,
            page_num            TEXT,
            interviewee         TEXT,
            interview_date      TEXT,
            interview_location  TEXT,
            content             TEXT,
            imported_at         TEXT
        )
    """)
    new_cols = [
        ("doc_type","TEXT"),("author","TEXT"),("pub_year","TEXT"),("publisher","TEXT"),
        ("chapter","TEXT"),("section","TEXT"),("page_num","TEXT"),
        ("interviewee","TEXT"),("interview_date","TEXT"),("interview_location","TEXT"),
    ]
    existing = {row[0] for row in conn.execute("DESCRIBE documents").fetchall()}
    for col_name, col_type in new_cols:
        if col_name not in existing:
            try:
                conn.execute(f"ALTER TABLE documents ADD COLUMN {col_name} {col_type}")
                logger.info(f"数据库升级：新增列 {col_name}")
            except Exception:
                pass
    conn.close()
    logger.info(f"数据库初始化/升级完成: {db_path}")


def get_record_count(db_path: str) -> int:
    conn = duckdb.connect(db_path)
    result = conn.execute("SELECT COUNT(*) FROM documents").fetchone()
    conn.close()
    return result[0] if result else 0


def check_file_already_imported(db_path: str, filename: str) -> bool:
    conn = duckdb.connect(db_path)
    result = conn.execute("SELECT COUNT(*) FROM documents WHERE source_file = ?", [filename]).fetchone()
    conn.close()
    return (result[0] > 0) if result else False


def get_records_count_for_files(db_path: str, filenames: list) -> int:
    """计算共享库中指定文件列表的总记录数"""
    if not filenames:
        return 0
    conn = duckdb.connect(db_path, read_only=True)
    placeholders = ", ".join(["?" for _ in filenames])
    result = conn.execute(
        f"SELECT COUNT(*) FROM documents WHERE source_file IN ({placeholders})",
        filenames,
    ).fetchone()
    conn.close()
    return result[0] if result else 0


def delete_source_file(db_path: str, filename: str) -> int:
    """从数据库中删除指定来源文件的所有记录，返回删除的条数"""
    conn = duckdb.connect(db_path)
    result = conn.execute(
        "SELECT COUNT(*) FROM documents WHERE source_file = ?", [filename]
    ).fetchone()
    count = result[0] if result else 0
    conn.execute("DELETE FROM documents WHERE source_file = ?", [filename])
    conn.close()
    logger.info(f"已删除来源文件 [{filename}] 共 {count} 条记录")
    return count


# ════════════════════════════════════════════════════════════
# 工具函数
# ════════════════════════════════════════════════════════════

def detect_encoding(filepath: str) -> str:
    with open(filepath, "rb") as f:
        raw = f.read(50000)
    result = chardet.detect(raw)
    encoding = result.get("encoding", "utf-8") or "utf-8"
    mapping = {"GB2312":"gbk","GBK":"gbk","UTF-8-SIG":"utf-8-sig","UTF-8":"utf-8"}
    return mapping.get(encoding.upper(), encoding)


ALL_FIELDS = [
    "source_file","file_type","doc_type","year","date","page","title",
    "author","pub_year","publisher","chapter","section","page_num",
    "interviewee","interview_date","interview_location","content","imported_at",
]

def _insert_records(conn, records: list):
    """批量插入记录（executemany 一次提交，避免逐行事务开销）。"""
    if not records:
        return
    cols = ", ".join(ALL_FIELDS)
    placeholders = ", ".join(["?"] * len(ALL_FIELDS))
    rows = [[rec.get(f, "") or "" for f in ALL_FIELDS] for rec in records]
    conn.executemany(f"INSERT INTO documents ({cols}) VALUES ({placeholders})", rows)


def _count_lines_fast(filepath: str) -> int:
    """按字节块统计换行数（仅用于进度估算，比逐行迭代快一个量级）。"""
    count = 0
    with open(filepath, "rb") as f:
        while True:
            buf = f.read(8 * 1024 * 1024)
            if not buf:
                break
            count += buf.count(b"\n")
    return count


def _make_base(meta: dict, filename: str, file_type: str) -> dict:
    return {
        "source_file":        filename,
        "file_type":          file_type,
        "doc_type":           meta.get("doc_type", ""),
        "title":              meta.get("title", filename),
        "author":             meta.get("author", ""),
        "pub_year":           meta.get("pub_year", ""),
        "publisher":          meta.get("publisher", ""),
        "interviewee":        meta.get("interviewee", ""),
        "interview_date":     meta.get("interview_date", ""),
        "interview_location": meta.get("interview_location", ""),
        "imported_at":        datetime.now().isoformat(),
    }


def _parse_chapter_section(title: str) -> tuple:
    chapter_match = re.match(r"^(第[一二三四五六七八九十百\d]+章)", title)
    if chapter_match:
        chapter = chapter_match.group(1)
        rest = title[len(chapter):].strip()
        section_match = re.match(r"^(第[一二三四五六七八九十百\d]+节)", rest)
        section = (section_match.group(1) + " " + rest[len(section_match.group(1)):].strip()
                   if section_match else rest)
        return chapter, section
    numeric_match = re.match(r"^(\d+)\.(\d+)", title)
    if numeric_match:
        return f"第{numeric_match.group(1)}章", title
    return title, ""


def _looks_like_chapter_title(line: str) -> bool:
    if not line or len(line) > 60:
        return False
    patterns = [
        r"^第[一二三四五六七八九十百\d]+[章节篇]",
        r"^\d+[\.\s]",
        r"^[一二三四五六七八九十]+[、\.]",
        r"^Chapter\s+\d+",
    ]
    return any(re.match(p, line) for p in patterns)


# ════════════════════════════════════════════════════════════
# 1. CSV 报纸
# ════════════════════════════════════════════════════════════

def ingest_csv(db_path: str, filepath: str, progress_callback=None) -> dict:
    logger.info(f"开始导入 CSV: {filepath}")
    encoding = detect_encoding(filepath)
    CHUNK_SIZE = 50_000
    total_imported = 0
    skipped = 0
    imported_at = datetime.now().isoformat()
    conn = duckdb.connect(db_path)

    try:
        total_rows = max(_count_lines_fast(filepath) - 1, 1)
        chunks = pd.read_csv(
            filepath, chunksize=CHUNK_SIZE, encoding=encoding,
            encoding_errors="replace", dtype=str, on_bad_lines="skip",
        )
        insert_cols = ", ".join(ALL_FIELDS)
        for chunk_idx, chunk in enumerate(chunks):
            chunk.columns = [c.strip() for c in chunk.columns]
            rename_map = {k: v for k, v in CSV_COLUMN_MAP.items() if k in chunk.columns}
            chunk = chunk.rename(columns=rename_map)
            chunk = chunk.fillna("")
            for col in ALL_FIELDS:
                if col not in chunk.columns:
                    chunk[col] = ""
            chunk["source_file"] = os.path.basename(filepath)
            chunk["file_type"] = "csv"
            chunk["doc_type"] = "newspaper"
            chunk["imported_at"] = imported_at

            before = len(chunk)
            chunk = chunk[chunk["content"].str.strip() != ""]
            skipped += before - len(chunk)
            if len(chunk) == 0:
                continue

            # DataFrame 整块插入（DuckDB 零拷贝读取，比逐行快两个数量级）
            batch = chunk[ALL_FIELDS].astype(str)
            conn.register("_csv_batch", batch)
            conn.execute(f"INSERT INTO documents ({insert_cols}) SELECT {insert_cols} FROM _csv_batch")
            conn.unregister("_csv_batch")

            total_imported += len(chunk)
            if progress_callback:
                progress_callback(min((chunk_idx+1)*CHUNK_SIZE/total_rows, 1.0), total_imported)
    except Exception as e:
        logger.error(f"CSV 导入失败: {e}", exc_info=True)
        raise
    finally:
        conn.close()

    logger.info(f"CSV 导入完成: {total_imported} 条")
    return {"total_imported": total_imported, "skipped": skipped, "encoding": encoding}


# ════════════════════════════════════════════════════════════
# 2. PDF（有书签/无书签自动识别）
# ════════════════════════════════════════════════════════════

def _extract_pdf_bookmarks(reader) -> list:
    results = []
    try:
        def _walk(outlines):
            for item in outlines:
                if isinstance(item, list):
                    _walk(item)
                else:
                    try:
                        page_idx = reader.get_destination_page_number(item)
                        results.append((item.title, page_idx))
                    except Exception:
                        pass
        _walk(reader.outline)
    except Exception as e:
        logger.warning(f"读取书签失败: {e}")
    return results


def ingest_pdf(db_path: str, filepath: str, meta: dict, progress_callback=None) -> dict:
    try:
        import pypdf
    except ImportError:
        raise ImportError("请先安装 pypdf: pip install pypdf")

    logger.info(f"开始导入 PDF: {filepath}")
    filename = os.path.basename(filepath)
    base = _make_base(meta, filename, "pdf")
    records = []

    with open(filepath, "rb") as f:
        reader = pypdf.PdfReader(f)
        total_pages = len(reader.pages)
        bookmarks = _extract_pdf_bookmarks(reader)

        if bookmarks:
            logger.info(f"PDF 有书签 {len(bookmarks)} 个")
            bookmarks.sort(key=lambda x: x[1])
            for i, (bm_title, start_page) in enumerate(bookmarks):
                end_page = bookmarks[i+1][1] if i+1 < len(bookmarks) else total_pages
                text_parts = []
                for p in range(start_page, min(end_page, total_pages)):
                    t = reader.pages[p].extract_text() or ""
                    text_parts.append(t.strip())
                full_text = "\n".join(t for t in text_parts if t)
                if len(full_text) < 20:
                    continue
                chapter, section = _parse_chapter_section(bm_title)
                records.append({**base, "chapter": chapter, "section": section,
                                "page_num": f"第{start_page+1}–{min(end_page,total_pages)}页",
                                "content": full_text})
                if progress_callback:
                    progress_callback((i+1)/len(bookmarks), len(records))
        else:
            logger.info("PDF 无书签，按页导入")
            current_chapter = ""
            for page_num in range(total_pages):
                text = (reader.pages[page_num].extract_text() or "").strip()
                if len(text) < 20:
                    continue
                first_line = text.split("\n")[0].strip()
                if _looks_like_chapter_title(first_line):
                    current_chapter = first_line
                records.append({**base, "chapter": current_chapter, "section": "",
                                "page_num": f"第{page_num+1}页", "content": text})
                if progress_callback:
                    progress_callback((page_num+1)/total_pages, len(records))

    conn = duckdb.connect(db_path)
    _insert_records(conn, records)
    conn.close()
    total_imported = len(records)
    logger.info(f"PDF 导入完成: {total_imported} 条")
    return {"total_imported": total_imported, "skipped": 0, "encoding": "binary"}


# ════════════════════════════════════════════════════════════
# 3. DOCX（按 Heading 样式划分章节）
# ════════════════════════════════════════════════════════════

def ingest_docx(db_path: str, filepath: str, meta: dict, progress_callback=None) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请先安装 python-docx: pip install python-docx")

    logger.info(f"开始导入 DOCX: {filepath}")
    filename = os.path.basename(filepath)
    base = _make_base(meta, filename, "docx")
    doc = Document(filepath)
    records = []

    sections = []
    current_h1 = ""
    current_h2 = ""
    buffer = []

    def flush_buf(h1, h2, buf):
        text = "\n".join(buf).strip()
        if len(text) >= 20:
            sections.append((h1, h2, text))

    HEADING1 = {"Heading 1", "标题 1", "标题1", "heading 1"}
    HEADING2 = {"Heading 2", "标题 2", "标题2", "heading 2"}
    HEADING3 = {"Heading 3", "标题 3", "标题3", "heading 3"}

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if style_name in HEADING1:
            flush_buf(current_h1, current_h2, buffer)
            current_h1, current_h2, buffer = text, "", []
        elif style_name in HEADING2:
            flush_buf(current_h1, current_h2, buffer)
            current_h2, buffer = text, []
        elif style_name in HEADING3:
            flush_buf(current_h1, current_h2, buffer)
            current_h2 = f"{current_h2} > {text}" if current_h2 else text
            buffer = []
        else:
            buffer.append(text)
    flush_buf(current_h1, current_h2, buffer)

    if not sections:
        logger.warning("DOCX 未检测到标题样式，退回按段落导入")
        all_text = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        for p in [x for x in all_text.split("\n\n") if len(x.strip()) >= 20]:
            records.append({**base, "chapter":"","section":"","page_num":"","content":p})
    else:
        for i, (h1, h2, text) in enumerate(sections):
            records.append({**base, "chapter":h1, "section":h2, "page_num":"", "content":text})
            if progress_callback:
                progress_callback((i+1)/len(sections), len(records))

    conn = duckdb.connect(db_path)
    _insert_records(conn, records)
    conn.close()
    total_imported = len(records)
    logger.info(f"DOCX 导入完成: {total_imported} 条")
    return {"total_imported": total_imported, "skipped": 0, "encoding": "utf-8"}


# ════════════════════════════════════════════════════════════
# 4. TXT（自动识别章节结构）
# ════════════════════════════════════════════════════════════

def ingest_txt(db_path: str, filepath: str, meta: dict, progress_callback=None) -> dict:
    logger.info(f"开始导入 TXT: {filepath}")
    encoding = detect_encoding(filepath)
    filename = os.path.basename(filepath)
    base = _make_base(meta, filename, "txt")

    with open(filepath, encoding=encoding, errors="replace") as f:
        lines = f.readlines()

    records = []
    current_chapter = ""
    current_section = ""
    buffer = []

    def flush(ch, sec, buf):
        text = "\n".join(buf).strip()
        if len(text) >= 20:
            records.append({**base, "chapter":ch, "section":sec, "page_num":"", "content":text})

    for line in lines:
        line_s = line.strip()
        if not line_s:
            continue
        if re.match(r"^第[一二三四五六七八九十百\d]+章", line_s) or re.match(r"^Chapter\s+\d+", line_s, re.I):
            flush(current_chapter, current_section, buffer)
            current_chapter, current_section, buffer = line_s, "", []
        elif re.match(r"^第[一二三四五六七八九十百\d]+节", line_s) or re.match(r"^\d+\.\d+\s+\S", line_s):
            flush(current_chapter, current_section, buffer)
            current_section, buffer = line_s, []
        elif re.match(r"^[一二三四五六七八九十]+[、．.]", line_s) and len(line_s) < 50:
            flush(current_chapter, current_section, buffer)
            current_section, buffer = line_s, []
        else:
            buffer.append(line_s)

    flush(current_chapter, current_section, buffer)

    if not records:
        with open(filepath, encoding=encoding, errors="replace") as f:
            raw = f.read()
        for p in [x.strip() for x in raw.split("\n\n") if len(x.strip()) >= 20]:
            records.append({**base, "chapter":"","section":"","page_num":"","content":p})

    conn = duckdb.connect(db_path)
    _insert_records(conn, records)
    conn.close()
    total_imported = len(records)
    logger.info(f"TXT 导入完成: {total_imported} 条")
    return {"total_imported": total_imported, "skipped": 0, "encoding": encoding}


# ════════════════════════════════════════════════════════════
# 5. EPUB
# ════════════════════════════════════════════════════════════

def ingest_epub(db_path: str, filepath: str, meta: dict, progress_callback=None) -> dict:
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("请先安装依赖: pip install ebooklib beautifulsoup4")

    logger.info(f"开始导入 EPUB: {filepath}")
    filename = os.path.basename(filepath)
    base = _make_base(meta, filename, "epub")
    book = epub.read_epub(filepath)
    records = []
    items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))

    for i, item in enumerate(items):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        heading = soup.find(["h1","h2","h3"])
        chapter_title = heading.get_text(strip=True) if heading else ""
        for tag in soup.find_all(["h1","h2","h3"]):
            tag.decompose()
        text = re.sub(r"\n{3,}", "\n\n", soup.get_text(separator="\n").strip())
        if len(text) < 20:
            continue
        chapter, section = _parse_chapter_section(chapter_title) if chapter_title else ("","")
        records.append({**base, "chapter": chapter or chapter_title, "section": section,
                        "page_num": f"章节{i+1}", "content": text})
        if progress_callback:
            progress_callback((i+1)/len(items), len(records))

    conn = duckdb.connect(db_path)
    _insert_records(conn, records)
    conn.close()
    total_imported = len(records)
    logger.info(f"EPUB 导入完成: {total_imported} 条")
    return {"total_imported": total_imported, "skipped": 0, "encoding": "utf-8"}


# ════════════════════════════════════════════════════════════
# 6. MOBI（转 EPUB 后处理）
# ════════════════════════════════════════════════════════════

def ingest_mobi(db_path: str, filepath: str, meta: dict, progress_callback=None) -> dict:
    import subprocess, tempfile, shutil
    if not shutil.which("ebook-convert"):
        raise RuntimeError(
            "导入 MOBI 需要安装 Calibre。\n"
            "请访问 https://calibre-ebook.com/ 下载安装后重启软件。"
        )
    with tempfile.TemporaryDirectory() as tmpdir:
        epub_path = os.path.join(tmpdir, "converted.epub")
        subprocess.run(["ebook-convert", filepath, epub_path], check=True, capture_output=True)
        return ingest_epub(db_path, epub_path, meta, progress_callback)
